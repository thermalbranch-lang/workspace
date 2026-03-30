from __future__ import annotations

import argparse
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader


WORKSPACE = Path(__file__).parent
INPUT_REPORT = WORKSPACE / "1 IAQ and moisture assessment report template V1.docx"
OUTPUT_REPORT = WORKSPACE / "1 IAQ and moisture assessment report template V1 - Summers filled.docx"


@dataclass
class SampleResult:
    sample_id: str
    location: str
    sample_type: str
    result_summary: str
    fungi: str


def set_cell_text(cell, text: str) -> None:
    cell.text = text


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def delete_table(table) -> None:
    element = table._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def delete_paragraphs_by_prefix(doc: Document, prefixes: list[str]) -> None:
    matched = True
    while matched:
        matched = False
        for paragraph in list(doc.paragraphs):
            if any(paragraph.text.startswith(prefix) for prefix in prefixes):
                delete_paragraph(paragraph)
                matched = True
                break


def delete_section_between(doc: Document, start_prefix: str, end_prefix: str) -> None:
    removing = False
    for paragraph in list(doc.paragraphs):
        if paragraph.text.startswith(start_prefix):
            removing = True
        if removing:
            if paragraph.text.startswith(end_prefix):
                break
            delete_paragraph(paragraph)



def fill_cover_table(table) -> None:
    values = [
        "Property: Summers Residence",
        "Client: Deborah Summers",
        "Report issued: March 30, 2026",
        "Date(s) of assessment: March 11, 2026",
        "Prepared by: Investigative Inspection Services, Inc. (IIS)",
        "Consultant: Stanley Yeskolski",
    ]
    for row, value in zip(table.rows, values):
        set_cell_text(row.cells[0], value)


def fill_client_information(table) -> None:
    set_cell_text(table.rows[1].cells[1], "Deborah Summers")
    set_cell_text(table.rows[1].cells[3], "dboothsummers@gmail.com")
    set_cell_text(table.rows[2].cells[1], "(804) 399-2014")
    set_cell_text(table.rows[2].cells[3], "N/A")
    address = "7306 Woodway Lane, Norfolk, VA 23505"
    set_cell_text(table.rows[3].cells[1], address)
    set_cell_text(table.rows[4].cells[1], address)


def fill_assessment_details(table) -> None:
    set_cell_text(table.rows[1].cells[1], "Stanley Yeskolski")
    set_cell_text(table.rows[2].cells[1], "March 11, 2026")
    set_cell_text(table.rows[3].cells[1], "N/A")


def fill_environment(table) -> None:
    for column in range(1, 5):
        set_cell_text(table.rows[1].cells[column], "N/A")
        set_cell_text(table.rows[4].cells[column], "N/A")
        set_cell_text(table.rows[6].cells[column], "N/A")

    set_cell_text(table.rows[2].cells[1], "N/A")
    set_cell_text(table.rows[2].cells[2], "N/A")
    set_cell_text(table.rows[2].cells[3], "N/A")
    set_cell_text(table.rows[2].cells[4], "N/A")

    set_cell_text(table.rows[3].cells[1], "N/A")
    set_cell_text(table.rows[3].cells[2], "N/A")
    set_cell_text(table.rows[3].cells[3], "N/A")
    set_cell_text(table.rows[3].cells[4], "N/A")

    set_cell_text(table.rows[5].cells[1], "Not reported")
    set_cell_text(table.rows[5].cells[2], "Not reported")
    set_cell_text(table.rows[5].cells[3], "N/A")
    set_cell_text(table.rows[5].cells[4], "N/A")


def fill_building_details(table) -> None:
    set_cell_text(table.rows[1].cells[3], "Approx. 1951 (approximately 75 years old)")
    set_cell_text(table.rows[1].cells[6], "Rear addition present; date not reported")
    set_cell_text(table.rows[2].cells[3], "4")
    set_cell_text(table.rows[2].cells[6], "2")
    set_cell_text(table.rows[3].cells[3], "N/A")
    set_cell_text(table.rows[3].cells[6], "Norfolk, Virginia")
    set_cell_text(table.rows[4].cells[3], "Single-family residence")
    set_cell_text(table.rows[5].cells[3], "2 Daikin systems and 1 Mitsubishi mini split; equipment located in the garage, attic, and wall-mounted locations")
    set_cell_text(table.rows[6].cells[2], "Grading generally slopes away from the home; drip and spray irrigation are present")
    set_cell_text(table.rows[7].cells[2], "Encapsulated crawl space with slab-supported rear addition")
    set_cell_text(table.rows[8].cells[4], "Primarily brick exterior with wood at the rear addition")
    set_cell_text(table.rows[9].cells[1], "Mixed roof assembly; client reported prior repair to a flat roof area")
    set_cell_text(table.rows[10].cells[4], "Hardwood flooring, plaster ceilings, wood paneling, and drywall finishes were observed or reported")
    set_cell_text(table.rows[11].cells[1], "Conditioned attic with a dehumidifier, insulated rafter bays, and no soffit or ridge ventilation")


def fill_equipment_table(table) -> None:
    for row_index in (4, 6, 10, 18):
        set_cell_text(table.rows[row_index].cells[2], "Yes")


def build_report() -> None:
    doc = Document(INPUT_REPORT)

    fill_cover_table(doc.tables[0])
    fill_client_information(doc.tables[2])
    fill_assessment_details(doc.tables[3])
    fill_environment(doc.tables[4])
    fill_building_details(doc.tables[5])
    fill_equipment_table(doc.tables[6])

    doc.paragraphs[15].text = (
        "Investigative Inspection Services (IIS) was contacted to perform a limited moisture and mold "
        "assessment at 7306 Woodway Lane, Norfolk, Virginia 23505. The client reported chronic humidity, "
        "moisture, HVAC-related concerns, visible fungal growth, bubbling paint, ceiling cracking, and "
        "health symptoms that reportedly improve when the family is away from the home. A site visit "
        "occurred on March 11, 2026."
    )
    doc.paragraphs[16].text = "The scope of inspection included the following:"
    doc.paragraphs[17].text = (
        "A visual assessment of the exterior, first-story living room area, attic, crawl space, and accessible "
        "HVAC components to identify moisture issues, fungal reservoirs, structural concerns related to "
        "moisture exposure, and areas appropriate for targeted tape-lift sampling."
    )
    doc.paragraphs[18].text = (
        "Assessment of accessible moisture conditions and representative fungal contamination in areas of concern."
    )
    doc.paragraphs[21].text = (
        "Report of findings from the assessment, including laboratory results of tape-lift samples collected "
        "during the site visit."
    )
    doc.paragraphs[23].text = "Laboratory report of tape-lift samples for direct examination for fungi."
    doc.paragraphs[24].text = ""
    doc.paragraphs[25].text = ""

    doc.paragraphs[46].text = (
        "Rusting steel lintels were observed above multiple windows, with associated mortar separation and "
        "localized brick displacement. The garage lintel exhibited the greatest movement, with an approximate "
        "1/4-inch separation at the left end. The right chimney flue had visible cracking, both chimneys had "
        "minimal mortar crown protection, and vegetation was growing near the caps. Corrosion staining at the "
        "copper flashing suggests incompatible ferrous fasteners may have been used. Crawl-space vents at the "
        "front left and left side are below grade, the front stoop trim was beginning to soften adjacent to the "
        "brick, and the wood-to-brick transition at the rear addition was not sealed, creating a potential air "
        "and moisture entry path."
    )
    doc.paragraphs[47].text = (
        "A qualified masonry or exterior repair contractor should evaluate and repair the corroded lintels, "
        "repoint or replace separated mortar joints, repair chimney crown and flue defects, remove vegetation, "
        "and correct any incompatible flashing fasteners. Below-grade vent exposure, the unsealed addition joint, "
        "and the soft front stoop trim should be addressed to reduce continued moisture entry and air infiltration."
    )

    doc.paragraphs[53].text = (
        "At the transition between the first-story living room and the dining room addition, the living room "
        "floor was approximately 4 to 5 mm lower than the adjoining slab-supported area. No active staining was "
        "observed at the flooring or walls in this location during the assessment; however, the separation where "
        "the slab and pier-and-beam systems meet creates a pathway for crawl-space air to enter the occupied "
        "space. Prior fungal growth had reportedly been observed on furniture in this area. A tape-lift sample "
        "taken at the first-story floor returned light Aspergillus/Penicillium with trace mycelium, which is "
        "consistent with settled fungal contamination or low-level growth influence near the floor separation."
    )
    doc.paragraphs[54].text = (
        "The floor separation should be evaluated by a qualified contractor and sealed after the underlying "
        "moisture and air-infiltration issues are addressed. Until permanent repairs are made, temporary sealing "
        "measures should be considered to reduce crawl-space air migration into the first-story living area. "
        "Affected furnishings and horizontal surfaces in this area should be cleaned or evaluated by a qualified "
        "mold remediation contractor following remediation of the crawl space and attic reservoirs."
    )

    doc.paragraphs[60].text = (
        "The attic had been converted to a conditioned space with a dehumidifier and did not have soffit or ridge "
        "ventilation. Wooden siding had been installed over the rafters and the rafter bays were insulated. "
        "Moisture readings taken through drilled access points did not show elevated moisture at the time of the "
        "assessment, and the openings were resealed with brown caulking. Even so, visible fungal growth was "
        "observed sporadically to heavily on the wood surfaces facing the attic air. Tape-lift analysis confirmed "
        "very heavy Penicillium with many mycelial fragments at the attic HVAC location and at the left attic "
        "ceiling sample, while the middle attic ceiling sample showed no fungi detected. The attic access also "
        "showed evidence of air leakage."
    )
    doc.paragraphs[61].text = (
        "The attic requires professional mold remediation, including cleaning or removal of contaminated "
        "materials as determined by the remediation contractor. The attic access should be improved with a gasket "
        "or other air-sealing method, and humidity control should be maintained after remediation. Items stored in "
        "the attic should not be moved into occupied areas until they have been evaluated and, if appropriate, "
        "cleaned as potentially contaminated contents."
    )

    doc.paragraphs[66].text = (
        "The crawl space was difficult to access because of ductwork, plumbing, HVAC lines, and abandoned hydronic "
        "piping. Relative humidity ranged from 59% to 66% at approximately 74 degrees Fahrenheit. Visible fungal "
        "growth was observed throughout the crawl space, with coverage varying from light to heavy. Laboratory "
        "results confirmed very heavy Cladosporium and heavy to very heavy Aspergillus/Penicillium on crawl-space "
        "joists, with mycelial estimates ranging from few to many. A condensate line at the rear center was leaking "
        "at a taped tubing-to-PVC connection, leaving a small puddle below. The main trunk supply ducts were not "
        "insulated, exposed fiberglass was visible at a flex-duct connection, and the vapor barrier attachment was "
        "incomplete or inadequately fastened in several areas. At the front left, the sill plate and floor-joist "
        "bearing ends were severely deteriorated from apparent historic water intrusion at the vents. Although the "
        "wood appeared dry at the time of the assessment, the damage is structural in nature. A cut floor joist was "
        "also noted under the guest bathroom."
    )
    doc.paragraphs[67].text = (
        "The crawl space requires professional mold remediation and corrective moisture control. The condensate leak "
        "should be repaired promptly, the crawl-space air barrier should be improved to reduce air migration into the "
        "home, and the vapor barrier should be properly secured where accessible. A qualified contractor should repair "
        "the deteriorated sill plate and joist bearing areas, evaluate the cut joist condition, and address below-grade "
        "vent exposure. Insulating the metal trunk ducts and improving dehumidifier air circulation, potentially through "
        "ducting or redistribution, would also be advisable."
    )

    doc.paragraphs[79].text = (
        "Evidence of fungal contamination was observed within the HVAC system. A tape-lift sample from inside the air "
        "handler identified rare Sporidesmium, Alternaria, and Curvularia, while a separate insulation sample from the "
        "HVAC system showed no fungi detected. By contrast, the attic-side HVAC sample returned very heavy Penicillium "
        "with many mycelial fragments, confirming that fungal growth is present on HVAC-adjacent surfaces in the attic. "
        "These findings indicate that at least one HVAC system serving the affected areas has been impacted by fungal "
        "contamination even though not every sampled component showed growth."
    )
    doc.paragraphs[80].text = (
        "The HVAC system serving the affected areas should be evaluated and restored by a qualified contractor familiar "
        "with mold-contaminated systems and the current NADCA standard. Contaminated components should be cleaned or "
        "replaced as appropriate after source removal has been completed in the attic and crawl space. The system should "
        "not be relied upon to restore indoor air quality until remediation and post-remediation cleaning are complete."
    )

    doc.paragraphs[86].text = (
        "Project: Summers Residence\n"
        "Lab: Hayes Microbial Consulting, LLC\n"
        "Report Number: 26011751\n"
        "Collection Date: March 11, 2026\n"
        "Analysis Date: March 12, 2026"
    )
    doc.paragraphs[87].text = (
        "Samples Collected: Nine tape-lift samples were collected from the HVAC system, the first-story floor area, "
        "the attic wood surfaces and attic-side HVAC location, and crawl-space framing."
    )
    doc.paragraphs[88].text = (
        "Critical Findings: Very heavy Penicillium was identified at the attic HVAC and attic ceiling left sample "
        "locations. Crawl-space joists returned very heavy Cladosporium and heavy to very heavy Aspergillus/Penicillium. "
        "The first-story floor sample returned light Aspergillus/Penicillium, while the HVAC insulation sample and attic "
        "ceiling middle sample both showed no fungi detected."
    )
    doc.paragraphs[94].text = "Tape-Lift Sample Results (Hayes Microbial Report 26011751)"
    doc.paragraphs[95].text = (
        "Sample 111, collected from inside the air handler, identified rare Sporidesmium, Alternaria, and Curvularia. "
        "Sample 222, collected from HVAC insulation, showed no fungi detected. Sample 333, collected at the first-story "
        "floor, identified light Aspergillus/Penicillium with trace mycelium."
    )
    doc.paragraphs[96].text = (
        "Sample B4373308 from the attic-side HVAC location identified very heavy Penicillium with many mycelial fragments. "
        "Sample B4326704 from the attic ceiling middle showed no fungi detected. Sample B4296879 from the left side of the "
        "attic ceiling identified very heavy Penicillium with many mycelial fragments."
    )
    doc.paragraphs[97].text = (
        "In the crawl space, sample B4326729 from a near-center joist identified very heavy Cladosporium with few mycelial "
        "fragments. Sample B4326696 from the center joist identified very heavy Aspergillus/Penicillium with many mycelial "
        "fragments, and sample B4343139 from the front center joist identified heavy Aspergillus/Penicillium with many "
        "mycelial fragments."
    )
    doc.paragraphs[98].text = (
        "The attic and crawl-space samples are consistent with active fungal growth rather than incidental spore deposition. "
        "The combination of heavy to very heavy spore loading and the presence of mycelial fragments demonstrates that these "
        "areas contain established fungal reservoirs associated with chronic moisture exposure or prolonged humidity."
    )
    doc.paragraphs[99].text = (
        "The light Aspergillus/Penicillium finding at the first-story floor supports the field observation that air and particulate "
        "migration from below the floor assembly is likely affecting the occupied space. Taken together with the attic and crawl-space "
        "results, the laboratory findings support the need for professional remediation, source control, and detailed cleaning of the "
        "affected living areas and HVAC components."
    )

    doc.paragraphs[144].text = (
        "The inspection identified fungal growth and moisture-related defects in the attic, crawl space, HVAC system, and first-story "
        "living area. The most significant findings were heavy fungal reservoirs in the crawl space and attic, active or recent HVAC-related "
        "contamination, and a floor separation at the first story that can allow contaminated crawl-space air to enter the living space."
    )
    doc.paragraphs[145].text = (
        "The crawl space contained widespread visible fungal growth and laboratory-confirmed heavy contamination on framing members, along with "
        "a leaking condensate connection, inadequate vapor-barrier attachment in areas, and structural deterioration at the front-left sill plate "
        "and joist bearing location. The attic also contained laboratory-confirmed heavy Penicillium growth, including at the attic-side HVAC location."
    )
    doc.paragraphs[146].text = (
        "The first-story living room showed floor separation where the slab-supported addition meets the pier-and-beam framing, and this condition, "
        "combined with prior fungal growth on furnishings and the light Aspergillus/Penicillium floor sample, supports the conclusion that fungal "
        "particulates are migrating into occupied areas."
    )
    doc.paragraphs[147].text = (
        "Immediate Action Required: Retain a qualified mold remediation contractor to remediate the attic, crawl space, contaminated HVAC components, "
        "and affected first-story contents or surfaces; repair the leaking condensate line; and implement temporary measures to reduce crawl-space air "
        "entry through the floor separation until permanent repairs are completed."
    )
    doc.paragraphs[148].text = (
        "Secondary Priority: Retain qualified contractors to repair the structurally deteriorated crawl-space framing, improve crawl-space air sealing and "
        "vapor-barrier attachment, evaluate the cut joist and ongoing floor movement, and address the exterior masonry, chimney, and flashing defects that "
        "may contribute to future moisture problems."
    )
    doc.paragraphs[149].text = (
        "The homeowner is strongly encouraged to address these items promptly to limit continued fungal amplification, reduce indoor exposure, and protect "
        "the structural components of the home."
    )
    doc.paragraphs[150].text = (
        "Based on the visible conditions and laboratory findings, Condition 3 exists in the attic, crawl space, and portions of the HVAC system. Condition 2 "
        "is also likely present in the first-story living area and on contents that have been exposed to migrated fungal particulates. After remediation, the "
        "home and affected HVAC components should be thoroughly cleaned and verified dry before rebuilding or returning stored contents to use."
    )
    doc.paragraphs[164].text = (
        "The property is occupied by individuals who reported coughing or throat irritation, nasal congestion or sneezing, sinus problems or headaches, and "
        "skin irritation that reportedly improve when away from the home. All repairs and remediation work should therefore be performed in strict compliance "
        "with current standards applicable to health-sensitive occupancy."
    )

    for table_index in [20, 19, 18, 17, 16, 15, 14, 13, 12, 11, 10, 9, 8, 7]:
        delete_table(doc.tables[table_index])

    delete_paragraphs_by_prefix(
        doc,
        [
            "Air samples are collected via a non-viable Allergenco-D cassette spore trap. This type of spore trap collects and traps fungal spores and fragments. A calibration is conducted before and after sampling to ensure a consistent rate of flow.",
            "Air sampling was done according to the following:",
            "Indoor Environmental Standards Organization (IESO) Standard 1210",
            "The American Conference of Governmental Industrial Hygienists (ACGIH) Bioaerosols Assessment and Control",
        ],
    )
    delete_section_between(doc, "Basement ", "HVAC systems")
    delete_section_between(doc, "Analysis of Air Samples", "Conclusion")

    doc.save(OUTPUT_REPORT)


def extract_docx_text(path: Path) -> Path:
    document = Document(path)
    output = path.with_suffix(path.suffix + ".txt")
    output.write_text("\n".join(paragraph.text for paragraph in document.paragraphs), encoding="utf-8")
    return output


def extract_pdf_texts() -> dict[str, str]:
    texts: dict[str, str] = {}
    for path in sorted((WORKSPACE / "Lab reports").glob("*.pdf")):
        reader = PdfReader(str(path))
        text = "\n".join((page.extract_text() or "") for page in reader.pages)
        out_path = Path(f"{path}.txt")
        out_path.write_text(text, encoding="utf-8")
        texts[path.name] = text
    return texts


def inspect_template() -> None:
    doc = Document(WORKSPACE / "1 IAQ and moisture assessment report template V1.docx")
    print(f"paragraphs {len(doc.paragraphs)}")
    print(f"tables {len(doc.tables)}")
    for index, paragraph in enumerate(doc.paragraphs):
        print(f"{index}: {paragraph.text}")
    for table_index, table in enumerate(doc.tables):
        print(f"TABLE {table_index} rows={len(table.rows)} cols={len(table.columns)}")
        for row_index, row in enumerate(table.rows):
            values = [cell.text.replace("\n", " | ") for cell in row.cells]
            print(f"ROW {row_index}: {values}")

    print("BLOCK ORDER")
    table_lookup = {id(table._tbl): index for index, table in enumerate(doc.tables)}
    para_counter = 0
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            paragraph = Paragraph(child, doc)
            print(f"P {para_counter}: {paragraph.text}")
            para_counter += 1
        elif isinstance(child, CT_Tbl):
            table = Table(child, doc)
            table_index = table_lookup[id(table._tbl)]
            print(f"T {table_index}: rows={len(table.rows)} cols={len(table.columns)}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--extract-pdfs", action="store_true")
    parser.add_argument("--inspect-template", action="store_true")
    parser.add_argument("--build-report", action="store_true")
    args = parser.parse_args()

    if args.extract_pdfs:
        extract_pdf_texts()
        print("Extracted lab PDF text files.")
    if args.inspect_template:
        inspect_template()
    if args.build_report:
        build_report()
        text_path = extract_docx_text(OUTPUT_REPORT)
        print(f"Built report: {OUTPUT_REPORT}")
        print(f"Extracted text: {text_path}")


if __name__ == "__main__":
    main()